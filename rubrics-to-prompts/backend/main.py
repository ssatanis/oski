import os
import io
import json
import yaml
import asyncio
from typing import List, Optional, Dict, Any
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import pytesseract
# Remove automatic EasyOCR import to handle SSL issues
try:
    import easyocr
    EASYOCR_AVAILABLE = True
except Exception as e:
    print(f"EasyOCR not available: {e}")
    EASYOCR_AVAILABLE = False
from PIL import Image
import aiofiles
from pydantic import BaseModel, ValidationError
from dotenv import load_dotenv
from openai import AzureOpenAI
from models import RubricPrompt, RubricSection, RubricItem
import tempfile
import logging
import PyPDF2
from docx import Document
import openpyxl

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Rubrics to Prompts API", version="1.0.0")

# Configure CORS
# For development, allow all origins. In production, restrict this.
if os.getenv("DEBUG", "False").lower() == "true":
    allowed_origins = ["*"]
else:
    allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(",")
    allowed_origins = [origin.strip() for origin in allowed_origins]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Azure OpenAI client
azure_openai_client = None
azure_openai_key = os.getenv("AZURE_OPENAI_KEY")
azure_openai_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
azure_openai_deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

if azure_openai_key and azure_openai_endpoint and azure_openai_key != "your_azure_openai_key_here":
    try:
        azure_openai_client = AzureOpenAI(
            api_version="2024-12-01-preview",
            azure_endpoint=azure_openai_endpoint,
            api_key=azure_openai_key
        )
        logger.info("Azure OpenAI client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Azure OpenAI client: {e}")
        azure_openai_client = None
else:
    logger.warning("Azure OpenAI configuration missing or using placeholder values. Using fallback processing.")

# Initialize EasyOCR reader (lazy loading with error handling)
ocr_reader = None

def get_ocr_reader():
    global ocr_reader
    if ocr_reader is None and EASYOCR_AVAILABLE:
        try:
            ocr_reader = easyocr.Reader(['en'])
            logger.info("EasyOCR initialized successfully")
        except Exception as e:
            logger.warning(f"Failed to initialize EasyOCR: {e}. Using Tesseract only.")
            ocr_reader = "failed"
    return ocr_reader if ocr_reader != "failed" else None

class ProcessingStatus(BaseModel):
    step: str
    message: str
    progress: float
    completed: bool = False

class ProcessingResponse(BaseModel):
    task_id: str
    status: str
    message: str

# In-memory task storage (use Redis in production)
task_storage = {}

async def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from various file formats using OCR"""
    try:
        if file_type.lower() in ['pdf', 'png', 'jpg', 'jpeg', 'tiff', 'bmp']:
            # Use Tesseract (primary) and EasyOCR (secondary if available)
            
            # Tesseract OCR
            try:
                if file_type.lower() == 'pdf':
                    # For PDF files, we can't directly open as image
                    # Skip Tesseract for PDF and rely on PDF text extraction
                    logger.info("PDF detected, skipping Tesseract OCR for image processing")
                    tesseract_text = ""
                else:
                    image = Image.open(file_path).convert('RGB')
                    tesseract_config = os.getenv('TESSERACT_CONFIG', '--psm 6')
                    tesseract_text = pytesseract.image_to_string(image, config=tesseract_config)
                    logger.info(f"Tesseract extracted {len(tesseract_text)} characters")
                    
                    # Return immediately if Tesseract got good results
                    if len(tesseract_text.strip()) > 50:
                        return tesseract_text
            except Exception as e:
                logger.warning(f"Tesseract OCR failed: {e}")
                tesseract_text = ""
            
            # EasyOCR (only if Tesseract didn't get good results and EasyOCR is available)
            easyocr_text = ""
            if len(tesseract_text.strip()) < 50:
                try:
                    reader = get_ocr_reader()
                    if reader and file_type.lower() != 'pdf':
                        easyocr_results = reader.readtext(file_path)
                        easyocr_text = ' '.join([result[1] for result in easyocr_results])
                        logger.info(f"EasyOCR extracted {len(easyocr_text)} characters")
                except Exception as e:
                    logger.warning(f"EasyOCR failed: {e}")
                    easyocr_text = ""
            
            # Combine results, preferring the longer extraction
            if len(easyocr_text) > len(tesseract_text):
                extracted_text = easyocr_text
                logger.info("Using EasyOCR result")
            elif len(tesseract_text) > 0:
                extracted_text = tesseract_text
                logger.info("Using Tesseract result")
            else:
                raise ValueError("OCR failed to extract meaningful text from image")
        
        elif file_type.lower() in ['txt', 'csv']:
            # Read text files directly with encoding fallback
            try:
                async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
                    extracted_text = await f.read()
            except UnicodeDecodeError:
                # Try with different encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        async with aiofiles.open(file_path, mode='r', encoding=encoding) as f:
                            extracted_text = await f.read()
                        logger.info(f"Successfully read file with {encoding} encoding")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ValueError("Could not decode text file with any supported encoding")
        
        elif file_type.lower() == 'pdf':
            # Extract text from PDF using PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    extracted_text = ""
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        extracted_text += page_text + "\n"
                    
                    # Check if we extracted meaningful text
                    if len(extracted_text.strip()) < 50:
                        raise ValueError("PDF appears to be image-based or contains no extractable text")
                    
                    logger.info(f"Extracted {len(extracted_text)} characters from PDF")
            except Exception as e:
                logger.warning(f"PDF text extraction failed: {e}")
                raise ValueError("Could not extract text from PDF. Please ensure it's a text-based PDF or convert to image format.")
        
        elif file_type.lower() in ['docx']:
            # Extract text from Word documents using python-docx
            try:
                doc = Document(file_path)
                extracted_text = ""
                
                for paragraph in doc.paragraphs:
                    extracted_text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            extracted_text += cell.text + "\t"
                        extracted_text += "\n"
                
                if len(extracted_text.strip()) < 10:
                    raise ValueError("No meaningful text found in Word document")
                
                logger.info(f"Extracted {len(extracted_text)} characters from Word document")
            except Exception as e:
                logger.warning(f"Word document processing failed: {e}")
                raise ValueError("Could not process Word document. Please save as PDF or text format.")
        
        elif file_type.lower() in ['xlsx', 'xls']:
            # Extract text from Excel files using openpyxl
            try:
                if file_type.lower() == 'xlsx':
                    workbook = openpyxl.load_workbook(file_path)
                    extracted_text = ""
                    
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        extracted_text += f"Sheet: {sheet_name}\n"
                        
                        for row in sheet.iter_rows(values_only=True):
                            row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                            if row_text.strip():
                                extracted_text += row_text + "\n"
                    
                    if len(extracted_text.strip()) < 10:
                        raise ValueError("No meaningful text found in Excel file")
                    
                    logger.info(f"Extracted {len(extracted_text)} characters from Excel file")
                else:
                    # For .xls files, we'd need xlrd library
                    raise ValueError("Legacy Excel (.xls) format not supported. Please save as .xlsx")
            except Exception as e:
                logger.warning(f"Excel file processing failed: {e}")
                raise ValueError("Could not process Excel file. Please save as PDF or text format.")
        
        elif file_type.lower() in ['doc']:
            # For legacy Word documents, provide helpful error
            raise ValueError("Legacy Word (.doc) format not supported. Please save as .docx, PDF, or text format.")
        
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return extracted_text.strip()
    
    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        raise HTTPException(status_code=400, detail=f"Failed to extract text from file: {str(e)}")

async def generate_yaml_prompt(rubric_text: str, task_id: str) -> Dict[str, Any]:
    """Generate YAML prompt using Azure OpenAI with retry logic and Pydantic validation"""
    
    # Update task status
    task_storage[task_id]['status'] = ProcessingStatus(
        step="ai_analysis",
        message="Analyzing rubric structure with AI...",
        progress=60.0
    )
    
    # Use the format from the example 1A.yaml
    prompt = f"""
You are an expert at converting OSCE exam rubrics into structured YAML prompts for AI assessment systems.

Convert the following rubric text into a YAML format that follows this EXACT structure:

key: 
  [station_identifier]
system_message: 
  |
   You are a helpful assistant tasked with analyzing and scoring a recorded medical examination between a medical student and a patient. Provide your response in JSON format.
   
user_message: 
  |   
   Important Instruction:
   When determining the start and end times of each examination, focus on the moments where the doctor instructs the patient to perform an action (e.g., "look up at the ceiling", "look straight ahead"). Give these phrases priority for setting the `start_time` and `end_time` over phrases where the doctor states their own actions (e.g., "I'm going to look at your nose and eyes").
      
   You need to identify the following physical exams from this conversation: 
   [List the specific examinations from the rubric here]
   
   If there is any part in the conversation where the medical student is listening to something but you cannot tell what specific organ it is, look at the conversation before and after to find what type of exam that was. Pay close attention to surrounding context and related physical examinations mentioned.
   
   If no exam is detected, you can say "No exam was performed", start_time: "nan", end_time: "nan", score: 0.
        
   # Formatting instructions
   
   - Ensure strict adherence to JSON formatting.
   - Do not use double quotes for multiple statements within a single field.
   - Use commas, single quotes, or other appropriate delimiters for multiple statements.
   - Do not include any text before or after the JSON output. Provide ONLY the json response.
   
   Please provide a response in the following format with keys: [list the exam keys]
   
   and the schema: 
   {{
        "statement": "statement extracted from the conversation that supports this specific exam",
        "start_time": "timepoint for start of the exam (ONLY 1 decimal pt)",
        "end_time": "timepoint for end of the exam (ONLY 1 decimal pt)",
        "rationale": "reasoning behind scoring the physical exam",
        "score": "score of the exam (0 or 1)"
   }}
response_config:
  structured_output: True

RUBRIC TEXT TO CONVERT:
{rubric_text}

Your task is to:
1. Extract the station identifier from the rubric (e.g., "1A", "2B", etc.)
2. Identify the physical examinations/procedures being assessed
3. Create the user_message that describes what to look for in the conversation
4. List the specific exams in the format shown above
5. Generate appropriate examples for each exam type

Generate ONLY the YAML content, without any markdown formatting, comments, or headers.
"""

    # Check if OpenAI is configured
    if azure_openai_client is None:
        logger.warning("OpenAI API not configured, using mock response")
        # Return a mock YAML for demo purposes
        mock_yaml = f"""key: 
  demo_station
system_message: 
  |
   You are a helpful assistant tasked with analyzing and scoring a recorded medical examination between a medical student and a patient. Provide your response in JSON format.
   
user_message: 
  |   
   You need to identify the following physical exams from this conversation: 
   1: Physical_Examination: Did the doctor perform a physical examination? 
    - Examples: 
         - I'm going to examine your abdomen
         - Let me listen to your heart
         - I'm going to check your reflexes
   
   Please provide a response in the following format with keys: Physical_Examination
   
   and the schema: 
   {{
        "statement": "statement extracted from the conversation that supports this specific exam",
        "start_time": "timepoint for start of the exam (ONLY 1 decimal pt)",
        "end_time": "timepoint for end of the exam (ONLY 1 decimal pt)",
        "rationale": "reasoning behind scoring the physical exam",
        "score": "score of the exam (0 or 1)"
   }}
response_config:
  structured_output: True"""
        
        return {
            "yaml_content": mock_yaml,
            "parsed_yaml": yaml.safe_load(mock_yaml),
            "validation_success": True
        }

    max_retries = 3
    for attempt in range(max_retries):
        try:
            # Update progress
            task_storage[task_id]['status'] = ProcessingStatus(
                step="ai_generation",
                message=f"AI processing attempt {attempt + 1}/{max_retries}...",
                progress=60.0 + (attempt * 10)
            )
            
            response = azure_openai_client.chat.completions.create(
                model=azure_openai_deployment or "2024-12-01-preview",
                messages=[
                    {"role": "system", "content": "You are an expert at converting medical assessment rubrics into structured YAML formats."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.3
            )
            
            yaml_content = response.choices[0].message.content.strip()
            
            # Clean up the response (remove markdown if present)
            if yaml_content.startswith("```yaml"):
                yaml_content = yaml_content[7:]
            if yaml_content.startswith("```"):
                yaml_content = yaml_content[3:]
            if yaml_content.endswith("```"):
                yaml_content = yaml_content[:-3]
            
            yaml_content = yaml_content.strip()
            
            # Validate YAML
            try:
                parsed_yaml = yaml.safe_load(yaml_content)
                logger.info("YAML validation successful")
                
                # Update task status
                task_storage[task_id]['status'] = ProcessingStatus(
                    step="validation",
                    message="Validating generated YAML...",
                    progress=90.0
                )
                
                return {
                    "yaml_content": yaml_content,
                    "parsed_yaml": parsed_yaml,
                    "validation_success": True
                }
            
            except yaml.YAMLError as e:
                logger.warning(f"YAML validation failed on attempt {attempt + 1}: {e}")
                if attempt == max_retries - 1:
                    raise HTTPException(status_code=500, detail=f"Generated YAML is invalid: {str(e)}")
                continue
        
        except Exception as e:
            logger.error(f"AI generation failed on attempt {attempt + 1}: {e}")
            if attempt == max_retries - 1:
                raise HTTPException(status_code=500, detail=f"Failed to generate YAML after {max_retries} attempts: {str(e)}")
            await asyncio.sleep(1)  # Brief delay before retry
    
    raise HTTPException(status_code=500, detail="Failed to generate valid YAML")

@app.post("/upload-rubric")
async def upload_rubric(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    """Upload and process a rubric file"""
    
    # Generate task ID
    import uuid
    task_id = str(uuid.uuid4())
    
    # Initialize task storage
    task_storage[task_id] = {
        'status': ProcessingStatus(
            step="upload",
            message="File uploaded successfully",
            progress=10.0
        ),
        'result': None,
        'error': None
    }
    
    # Validate file type
    allowed_extensions = ['pdf', 'xls', 'doc', 'docx', 'txt', 'xlsx', 'csv', 'png', 'jpg', 'jpeg']
    file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
    
    if file_extension not in allowed_extensions:
        task_storage[task_id]['error'] = f"Unsupported file type: {file_extension}"
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}")
    
    # Read file content before background processing
    file_content = await file.read()
    
    # Start background processing
    background_tasks.add_task(process_rubric_background, file_content, file.filename, task_id, file_extension)
    
    return ProcessingResponse(
        task_id=task_id,
        status="processing",
        message="File uploaded successfully. Processing started."
    )

async def process_rubric_background(file_content: bytes, filename: str, task_id: str, file_extension: str):
    """Background task to process the rubric file"""
    try:
        # Update status: File processing
        task_storage[task_id]['status'] = ProcessingStatus(
            step="file_processing",
            message="Processing uploaded file...",
            progress=20.0
        )
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Update status: Text extraction
            task_storage[task_id]['status'] = ProcessingStatus(
                step="ocr_processing",
                message="Extracting text from document...",
                progress=40.0
            )
            
            # Extract text
            extracted_text = await extract_text_from_file(temp_file_path, file_extension)
            
            if not extracted_text.strip():
                raise HTTPException(status_code=400, detail="No text could be extracted from the file")
            
            logger.info(f"Extracted text length: {len(extracted_text)} characters")
            
            # Generate YAML prompt
            result = await generate_yaml_prompt(extracted_text, task_id)
            
            # Update final status
            task_storage[task_id]['status'] = ProcessingStatus(
                step="completed",
                message="Processing completed successfully!",
                progress=100.0,
                completed=True
            )
            
            # Store result
            task_storage[task_id]['result'] = {
                'original_text': extracted_text,
                'yaml_content': result['yaml_content'],
                'parsed_yaml': result['parsed_yaml'],
                'filename': filename
            }
            
        finally:
            # Cleanup temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass
    
    except Exception as e:
        logger.error(f"Background processing failed: {e}")
        error_message = f"Load failed: {str(e)}"
        task_storage[task_id]['error'] = error_message
        task_storage[task_id]['status'] = ProcessingStatus(
            step="error",
            message=error_message,
            progress=0.0
        )

@app.get("/status/{task_id}")
async def get_processing_status(task_id: str):
    """Get the processing status of a task"""
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task_data = task_storage[task_id]
    
    if task_data.get('error'):
        return {
            "status": "error",
            "error": task_data['error'],
            "step": "error",
            "message": task_data['error'],
            "progress": 0.0
        }
    
    status = task_data['status']
    response = {
        "status": "completed" if status.completed else "processing",
        "step": status.step,
        "message": status.message,
        "progress": status.progress
    }
    
    if status.completed and task_data.get('result'):
        response['result'] = task_data['result']
    
    return response

@app.post("/update-yaml/{task_id}")
async def update_yaml(task_id: str, updated_yaml: Dict[str, Any]):
    """Update the generated YAML content"""
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task_data = task_storage[task_id]
    if not task_data.get('result'):
        raise HTTPException(status_code=400, detail="No result found to update")
    
    try:
        # Validate the updated YAML
        yaml_content = updated_yaml.get('yaml_content', '')
        parsed_yaml = yaml.safe_load(yaml_content)
        
        # Update the stored result
        task_data['result']['yaml_content'] = yaml_content
        task_data['result']['parsed_yaml'] = parsed_yaml
        
        return {"message": "YAML updated successfully", "status": "success"}
    
    except yaml.YAMLError as e:
        raise HTTPException(status_code=400, detail=f"Invalid YAML format: {str(e)}")

@app.get("/download-yaml/{task_id}")
async def download_yaml(task_id: str):
    """Download the generated YAML file"""
    if task_id not in task_storage:
        raise HTTPException(status_code=404, detail="Task not found")
    
    task_data = task_storage[task_id]
    result = task_data.get('result')
    
    if not result:
        raise HTTPException(status_code=400, detail="No result found to download")
    
    # Create temporary YAML file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.yaml', delete=False) as temp_file:
        temp_file.write(result['yaml_content'])
        temp_file_path = temp_file.name
    
    # Generate filename
    original_filename = result.get('filename', 'rubric')
    yaml_filename = f"{original_filename.split('.')[0]}_prompt.yaml"
    
    return FileResponse(
        path=temp_file_path,
        filename=yaml_filename,
        media_type='application/x-yaml'
    )

@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    """Simple upload endpoint that returns extracted text immediately"""
    try:
        # Validate file type
        allowed_extensions = ['pdf', 'xls', 'doc', 'docx', 'txt', 'xlsx', 'csv', 'png', 'jpg', 'jpeg']
        file_extension = file.filename.split('.')[-1].lower() if '.' in file.filename else ''
        
        if file_extension not in allowed_extensions:
            raise HTTPException(status_code=400, detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}")
        
        # Read file content
        file_content = await file.read()
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
            temp_file.write(file_content)
            temp_file_path = temp_file.name
        
        try:
            # Extract text
            extracted_text = await extract_text_from_file(temp_file_path, file_extension)
            
            if not extracted_text.strip():
                raise HTTPException(status_code=400, detail="No text could be extracted from the file")
            
            logger.info(f"Extracted text length: {len(extracted_text)} characters")
            
            return {
                "message": "File processed successfully",
                "filename": file.filename,
                "extracted_text": extracted_text,
                "text_length": len(extracted_text)
            }
            
        finally:
            # Cleanup temporary file
            try:
                os.unlink(temp_file_path)
            except:
                pass
                
    except Exception as e:
        logger.error(f"Upload processing failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")

class GeneratePromptRequest(BaseModel):
    extracted_text: str

@app.post("/generate-prompt")
async def generate_prompt(request: GeneratePromptRequest):
    """Generate prompt endpoint that takes extracted text and returns YAML/JSON"""
    try:
        # Generate a temporary task ID for tracking
        import uuid
        temp_task_id = str(uuid.uuid4())
        
        # Initialize temporary task storage
        task_storage[temp_task_id] = {
            'status': ProcessingStatus(
                step="ai_analysis",
                message="Analyzing rubric structure with AI...",
                progress=50.0
            ),
            'result': None,
            'error': None
        }
        
        # Generate YAML prompt
        result = await generate_yaml_prompt(request.extracted_text, temp_task_id)
        
        # Clean up temporary task
        if temp_task_id in task_storage:
            del task_storage[temp_task_id]
        
        # Convert YAML to JSON for the JSON download option
        import json
        json_content = json.dumps(result['parsed_yaml'], indent=2)
        
        return {
            "message": "Prompt generated successfully",
            "yaml_content": result['yaml_content'],
            "json_content": json_content,
            "parsed_yaml": result['parsed_yaml'],
            "validation_success": result['validation_success']
        }
        
    except Exception as e:
        logger.error(f"Prompt generation failed: {e}")
        # Clean up temporary task if it exists
        if 'temp_task_id' in locals() and temp_task_id in task_storage:
            del task_storage[temp_task_id]
        raise HTTPException(status_code=500, detail=f"Failed to generate prompt: {str(e)}")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "Rubrics to Prompts API"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000) 